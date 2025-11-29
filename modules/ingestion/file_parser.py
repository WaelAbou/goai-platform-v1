"""
File Parser Module - Page-Aware Document Extraction

Supports:
- PDF (with page numbers)
- DOCX (with page estimation)
- TXT/MD (with line-based pages)
- XLSX (with sheet names as pages)
"""

import io
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import hashlib
from datetime import datetime


@dataclass
class PageContent:
    """Represents content from a single page."""
    page_number: int
    content: str
    metadata: Dict[str, Any]


@dataclass
class ParsedDocument:
    """Represents a fully parsed document with page information."""
    filename: str
    file_type: str
    total_pages: int
    pages: List[PageContent]
    metadata: Dict[str, Any]
    document_id: str
    
    def get_full_text(self) -> str:
        """Get complete document text."""
        return "\n\n".join([p.content for p in self.pages])
    
    def get_page_text(self, page_number: int) -> Optional[str]:
        """Get text from a specific page."""
        for page in self.pages:
            if page.page_number == page_number:
                return page.content
        return None


class FileParser:
    """
    Page-aware file parser for various document formats.
    Extracts text while preserving page/section information.
    """
    
    SUPPORTED_TYPES = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'text': ['.txt', '.md', '.markdown', '.rst'],
        'excel': ['.xlsx', '.xls'],
        'csv': ['.csv'],
    }
    
    # Approximate characters per page for estimation
    CHARS_PER_PAGE = 3000
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which parsing libraries are available."""
        self.has_pypdf = False
        self.has_docx = False
        self.has_openpyxl = False
        
        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            pass
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass
        
        try:
            import openpyxl
            self.has_openpyxl = True
        except ImportError:
            pass
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """Determine file type from filename."""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if ext in extensions:
                return file_type
        return None
    
    def parse(
        self, 
        file_content: bytes, 
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ParsedDocument:
        """
        Parse a file and extract page-aware content.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            metadata: Optional additional metadata
            
        Returns:
            ParsedDocument with page information
        """
        file_type = self.get_file_type(filename)
        
        if file_type is None:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Generate document ID
        doc_id = hashlib.md5(
            f"{filename}_{datetime.now().isoformat()}".encode()
        ).hexdigest()[:16]
        
        # Parse based on file type
        if file_type == 'pdf':
            pages = self._parse_pdf(file_content, filename)
        elif file_type == 'docx':
            pages = self._parse_docx(file_content, filename)
        elif file_type == 'text':
            pages = self._parse_text(file_content, filename)
        elif file_type == 'excel':
            pages = self._parse_excel(file_content, filename)
        elif file_type == 'csv':
            pages = self._parse_csv(file_content, filename)
        else:
            raise ValueError(f"Parser not implemented for: {file_type}")
        
        # Build document metadata
        doc_metadata = {
            "filename": filename,
            "file_type": file_type,
            "file_size": len(file_content),
            "parsed_at": datetime.now().isoformat(),
            **(metadata or {})
        }
        
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            total_pages=len(pages),
            pages=pages,
            metadata=doc_metadata,
            document_id=doc_id
        )
    
    def _parse_pdf(self, content: bytes, filename: str) -> List[PageContent]:
        """Parse PDF with page-by-page extraction."""
        if not self.has_pypdf:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        from pypdf import PdfReader
        
        pages = []
        reader = PdfReader(io.BytesIO(content))
        
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            
            # Clean up text
            text = self._clean_text(text)
            
            if text.strip():  # Only add non-empty pages
                pages.append(PageContent(
                    page_number=i,
                    content=text,
                    metadata={
                        "source": filename,
                        "page": i,
                        "total_pages": len(reader.pages)
                    }
                ))
        
        return pages
    
    def _parse_docx(self, content: bytes, filename: str) -> List[PageContent]:
        """Parse DOCX with page estimation based on content length."""
        if not self.has_docx:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        
        # Extract all paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)
        
        combined_text = "\n".join(full_text)
        
        # Estimate pages based on character count
        pages = self._split_into_pages(combined_text, filename)
        
        return pages
    
    def _parse_text(self, content: bytes, filename: str) -> List[PageContent]:
        """Parse plain text files with page estimation."""
        # Detect encoding
        try:
            import chardet
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        except ImportError:
            encoding = 'utf-8'
        
        try:
            text = content.decode(encoding)
        except UnicodeDecodeError:
            text = content.decode('utf-8', errors='ignore')
        
        # Check for explicit page breaks
        if '\f' in text:  # Form feed character (page break)
            page_texts = text.split('\f')
            pages = []
            for i, page_text in enumerate(page_texts, 1):
                if page_text.strip():
                    pages.append(PageContent(
                        page_number=i,
                        content=self._clean_text(page_text),
                        metadata={"source": filename, "page": i}
                    ))
            return pages
        
        # Otherwise estimate pages
        return self._split_into_pages(text, filename)
    
    def _parse_excel(self, content: bytes, filename: str) -> List[PageContent]:
        """Parse Excel with each sheet as a page."""
        if not self.has_openpyxl:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")
        
        from openpyxl import load_workbook
        
        wb = load_workbook(io.BytesIO(content), data_only=True)
        pages = []
        
        for i, sheet_name in enumerate(wb.sheetnames, 1):
            sheet = wb[sheet_name]
            rows = []
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip() and row_text.replace("|", "").strip():
                    rows.append(row_text)
            
            if rows:
                pages.append(PageContent(
                    page_number=i,
                    content=f"Sheet: {sheet_name}\n\n" + "\n".join(rows),
                    metadata={
                        "source": filename,
                        "page": i,
                        "sheet_name": sheet_name
                    }
                ))
        
        return pages
    
    def _parse_csv(self, content: bytes, filename: str) -> List[PageContent]:
        """Parse CSV as a single page."""
        try:
            import chardet
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        except ImportError:
            encoding = 'utf-8'
        
        try:
            text = content.decode(encoding)
        except UnicodeDecodeError:
            text = content.decode('utf-8', errors='ignore')
        
        # Split into estimated pages if large
        return self._split_into_pages(text, filename)
    
    def _split_into_pages(
        self, 
        text: str, 
        filename: str,
        chars_per_page: Optional[int] = None
    ) -> List[PageContent]:
        """Split text into estimated pages based on character count."""
        chars_per_page = chars_per_page or self.CHARS_PER_PAGE
        text = self._clean_text(text)
        
        if len(text) <= chars_per_page:
            return [PageContent(
                page_number=1,
                content=text,
                metadata={"source": filename, "page": 1, "estimated": True}
            )]
        
        pages = []
        page_num = 1
        
        # Split by paragraphs to avoid cutting mid-sentence
        paragraphs = text.split('\n\n')
        current_page = []
        current_length = 0
        
        for para in paragraphs:
            para_len = len(para)
            
            if current_length + para_len > chars_per_page and current_page:
                # Save current page
                pages.append(PageContent(
                    page_number=page_num,
                    content="\n\n".join(current_page),
                    metadata={"source": filename, "page": page_num, "estimated": True}
                ))
                page_num += 1
                current_page = [para]
                current_length = para_len
            else:
                current_page.append(para)
                current_length += para_len
        
        # Don't forget the last page
        if current_page:
            pages.append(PageContent(
                page_number=page_num,
                content="\n\n".join(current_page),
                metadata={"source": filename, "page": page_num, "estimated": True}
            ))
        
        return pages
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Restore paragraph breaks
        text = re.sub(r'\s*\n\s*', '\n', text)
        # Remove control characters except newlines
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of all supported file extensions."""
        extensions = []
        for exts in self.SUPPORTED_TYPES.values():
            extensions.extend(exts)
        return extensions


# Singleton instance
file_parser = FileParser()

